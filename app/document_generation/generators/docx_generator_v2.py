import io
import os
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.document_generation.generators.base import BaseGenerator
from app.document_generation.markdown_ast import parse


class DOCXGeneratorV2(BaseGenerator):
    @property
    def format(self) -> str:
        return "docx-v2"

    @property
    def mime_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    @property
    def file_extension(self) -> str:
        return ".docx"

    def generate(self, html: str, template: Optional[Dict[str, Any]] = None, metadata: Optional[Dict[str, Any]] = None) -> bytes:
        return self._generate_v2(html, template, metadata)

    def generate_preview(self, html: str, template: Optional[Dict[str, Any]] = None) -> bytes:
        return self.generate(html, template)

    def _generate_v2(self, html: str, template: Optional[Dict] = None, metadata: Optional[Dict] = None) -> bytes:
        doc = Document()
        self._apply_template_styles(doc, template)
        ast = parse(html)
        self._walk_blocks(doc, ast, template)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _apply_template_styles(self, doc: Document, template: Optional[Dict] = None):
        if not template:
            return
        margins = template.get("margins", {})
        for section in doc.sections:
            if margins:
                section.top_margin = Inches(margins.get("top", 72) / 72)
                section.bottom_margin = Inches(margins.get("bottom", 72) / 72)
                section.left_margin = Inches(margins.get("left", 72) / 72)
                section.right_margin = Inches(margins.get("right", 72) / 72)

    def _walk_blocks(self, doc: Document, blocks: list, template: Optional[Dict] = None):
        for block in blocks:
            self._walk_block(doc, block, template)

    def _walk_block(self, doc: Document, block: dict, template: Optional[Dict] = None):
        t = block["type"]
        if t == "heading":
            self._add_heading(doc, block)
        elif t == "paragraph":
            self._add_paragraph(doc, block)
        elif t == "list":
            self._add_list(doc, block, depth=0)
        elif t == "block_code":
            self._add_code_block(doc, block)
        elif t == "block_quote":
            self._add_blockquote(doc, block)
        elif t == "table":
            self._add_table(doc, block)
        elif t == "thematic_break":
            doc.add_paragraph("_" * 60)
        elif t == "image":
            self._add_image(doc, block)

    def _add_heading(self, doc: Document, block: dict):
        text = self._render_inline(block.get("children", []))
        level = min(block.get("level", 1), 6)
        doc.add_heading(text, level=level)

    def _add_paragraph(self, doc: Document, block: dict):
        children = block.get("children", [])
        if not children:
            return
        p = doc.add_paragraph()
        self._render_inline_to_paragraph(p, children)

    def _add_list(self, doc: Document, block: dict, depth: int = 0):
        list_style = f"List Bullet" if not block.get("ordered") else "List Number"
        if depth > 0:
            list_style = f"List Bullet {depth + 1}" if depth < 9 else "List Bullet"
        for item in block.get("children", []):
            item_children = item.get("children", [])
            text_parts = []
            sub_list = None
            for child in item_children:
                if child["type"] == "list":
                    sub_list = child
                else:
                    text_parts.append(child)
            if text_parts:
                p = doc.add_paragraph(style=list_style)
                self._render_inline_to_paragraph(p, text_parts)
            else:
                p = doc.add_paragraph(style=list_style)
            if sub_list:
                self._add_list(doc, sub_list, depth + 1)

    def _add_code_block(self, doc: Document, block: dict):
        code_lines = block.get("raw", "").split("\n")
        for line in code_lines:
            if not line.strip() and len(code_lines) > 1:
                continue
            p = doc.add_paragraph()
            run = p.add_run(line if line.strip() else " ")
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

    def _add_blockquote(self, doc: Document, block: dict):
        for child in block.get("children", []):
            if child["type"] == "paragraph":
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                self._render_inline_to_paragraph(p, child.get("children", []), is_quote=True)

    def _add_table(self, doc: Document, block: dict):
        head_rows = []
        body_rows = []
        for child in block.get("children", []):
            if child["type"] == "table_head":
                for row in child.get("children", []):
                    head_rows.append(row)
            elif child["type"] == "table_body":
                for row in child.get("children", []):
                    body_rows.append(row)

        all_rows = head_rows + body_rows
        if not all_rows:
            return

        num_cols = 0
        for row in all_rows:
            num_cols = max(num_cols, len(row.get("children", [])))

        table = doc.add_table(rows=len(all_rows), cols=num_cols)
        table.style = "Table Grid"

        for ri, row_token in enumerate(all_rows):
            cells = row_token.get("children", [])
            is_header = ri < len(head_rows)
            for ci, cell_token in enumerate(cells):
                if ci >= num_cols:
                    break
                cell = table.cell(ri, ci)
                text = self._render_inline(cell_token.get("children", []))
                cell.text = text
                if is_header:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    def _add_image(self, doc: Document, block: dict):
        src = block.get("src", "")
        alt = block.get("alt", "")
        if src and not src.startswith("data:"):
            try:
                doc.add_picture(src, width=Inches(4))
            except Exception:
                p = doc.add_paragraph()
                p.add_run(f"[Image: {alt}]").italic = True

    def _render_inline(self, children: list) -> str:
        parts = []
        for child in children:
            t = child["type"]
            if t == "text":
                parts.append(child.get("raw", ""))
            elif t == "codespan":
                parts.append(child.get("raw", ""))
            elif t == "strong":
                parts.append(self._render_inline(child.get("children", [])))
            elif t == "emphasis":
                parts.append(self._render_inline(child.get("children", [])))
            elif t == "link":
                link_text = self._render_inline(child.get("children", []))
                url = child.get("url", "")
                parts.append(f"{link_text} ({url})" if url and link_text else link_text)
            elif t == "softbreak":
                parts.append(" ")
            elif t == "linebreak":
                parts.append("\n")
        return "".join(parts)

    def _render_inline_to_paragraph(self, p, children: list, is_quote: bool = False):
        colors = {"quote": RGBColor(0x66, 0x66, 0x66)}
        for child in children:
            self._render_inline_to_run(p, child, is_quote=is_quote)
        if is_quote:
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = colors["quote"]

    def _render_inline_to_run(self, p, child: dict, is_quote: bool = False, link_url: str = ""):
        t = child["type"]
        if t == "text":
            run = p.add_run(child.get("raw", ""))
        elif t == "codespan":
            run = p.add_run(child.get("raw", ""))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        elif t == "strong":
            for sc in child.get("children", []):
                self._render_inline_to_run(p, sc, is_quote=is_quote, link_url=link_url)
                if p.runs:
                    p.runs[-1].bold = True
        elif t == "emphasis":
            for sc in child.get("children", []):
                self._render_inline_to_run(p, sc, is_quote=is_quote, link_url=link_url)
                if p.runs:
                    p.runs[-1].italic = True
        elif t == "link":
            url = child.get("url", "")
            for sc in child.get("children", []):
                self._render_inline_to_run(p, sc, is_quote=is_quote, link_url=url)
                if p.runs:
                    self._add_hyperlink(p, p.runs[-1].text, url)
                    p._p.remove(p.runs[-1]._element)
        elif t == "softbreak":
            p.add_run(" ")
        elif t == "linebreak":
            p.add_run("\n")

    def _add_hyperlink(self, paragraph, text: str, url: str):
        from docx.oxml import OxmlElement
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        c = OxmlElement("w:color")
        c.set(qn("w:val"), "0563C1")
        rPr.append(c)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)


def docx_generate_v2(html: str, template: Optional[Dict] = None, metadata: Optional[Dict] = None) -> bytes:
    gen = DOCXGeneratorV2()
    return gen.generate(html, template, metadata)
