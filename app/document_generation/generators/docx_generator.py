import io
import re
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from html.parser import HTMLParser

from app.document_generation.generators.base import BaseGenerator


class _DocxHTMLParser(HTMLParser):
    def __init__(self, doc: Document):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.current_list_type = None
        self.current_list_index = 0
        self.in_header = False
        self.header_level = 0
        self.in_table = False
        self.in_table_row = False
        self.in_table_header = False
        self.table_data = []
        self.current_row = []
        self.in_code = False
        self.code_text = ""
        self.in_pre = False
        self.in_bold = False
        self.in_italic = False
        self.in_link = False
        self.link_href = ""
        self.in_paragraph = False

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        attrs_dict = dict(attrs)

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._end_paragraph()
            self.in_header = True
            self.header_level = int(tag[1])
            self.current_paragraph = self.doc.add_paragraph()
            style_name = f"Heading {self.header_level}"
            try:
                self.current_paragraph.style = self.doc.styles[style_name]
            except Exception:
                run = self.current_paragraph.add_run()
                run.bold = True
                font_sizes = {1: 24, 2: 20, 3: 16, 4: 14, 5: 12, 6: 11}
                run.font.size = Pt(font_sizes.get(self.header_level, 12))
            self.in_paragraph = True

        elif tag == "p":
            self._end_paragraph()
            self.current_paragraph = self.doc.add_paragraph()
            self.in_paragraph = True

        elif tag in ("ul", "ol"):
            self.current_list_type = tag
            self.current_list_index = 0

        elif tag == "li":
            self._end_paragraph()
            self.current_list_index += 1
            prefix = ""
            if self.current_list_type == "ol":
                prefix = f"{self.current_list_index}. "
            elif self.current_list_type == "ul":
                prefix = "• "
            self.current_paragraph = self.doc.add_paragraph(style="List Bullet")
            self.current_paragraph.paragraph_format.left_indent = Inches(0.5)
            self.in_paragraph = True

        elif tag == "strong" or tag == "b":
            self.in_bold = True
        elif tag == "em" or tag == "i":
            self.in_italic = True
        elif tag == "a":
            self.in_link = True
            self.link_href = attrs_dict.get("href", "")
        elif tag == "br":
            if self.current_paragraph:
                self.current_paragraph.add_run("\n")
        elif tag == "hr":
            self._end_paragraph()
            p = self.doc.add_paragraph()
            run = p.add_run("_" * 60)
            run.font.color.rgb = RGBColor(200, 200, 200)
        elif tag == "table":
            self.in_table = True
            self.table_data = []
        elif tag == "tr":
            self.current_row = []
        elif tag == "th":
            self.in_table_header = True
        elif tag == "td":
            self.in_table_header = False
        elif tag in ("pre", "code"):
            self.in_code = True
            self.code_text = ""

        elif tag == "blockquote":
            self._end_paragraph()
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.paragraph_format.left_indent = Inches(0.5)
            run = self.current_paragraph.add_run()
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            self.in_paragraph = True

        elif tag == "img":
            src = attrs_dict.get("src", "")
            alt = attrs_dict.get("alt", "")
            if src and not src.startswith("data:"):
                try:
                    self.doc.add_picture(src, width=Inches(4))
                except Exception:
                    if self.current_paragraph:
                        self.current_paragraph.add_run(f"[Image: {alt}]")

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self.in_header = False
        elif tag == "p":
            self._end_paragraph()
        elif tag == "li":
            self._end_paragraph()
        elif tag in ("strong", "b"):
            self.in_bold = False
        elif tag in ("em", "i"):
            self.in_italic = False
        elif tag == "a":
            self.in_link = False
        elif tag in ("pre", "code"):
            if self.code_text:
                if self.current_paragraph:
                    run = self.current_paragraph.add_run()
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
                    run.text = self.code_text
                else:
                    p = self.doc.add_paragraph()
                    run = p.add_run()
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
                    run.text = self.code_text
            self.code_text = ""
            self.in_code = False
        elif tag == "tr":
            if self.current_row:
                self.table_data.append(self.current_row)
                self.current_row = []
        elif tag == "table":
            self._render_table()
            self.in_table = False
        elif tag == "th":
            self.in_table_header = False
        elif tag == "blockquote":
            self._end_paragraph()
        elif tag == "ul":
            self.current_list_type = None
        elif tag == "ol":
            self.current_list_type = None

    def handle_data(self, data):
        if self.in_code:
            self.code_text += data
            return

        if self.in_table:
            self.current_row.append(data)
            return

        if self.current_paragraph:
            if self.in_link and self.link_href:
                self._add_hyperlink(self.current_paragraph, data, self.link_href)
            else:
                run = self.current_paragraph.add_run(data)
                if self.in_bold:
                    run.bold = True
                if self.in_italic:
                    run.italic = True
                if self.in_header:
                    run.bold = True

    def _end_paragraph(self):
        self.current_paragraph = None
        self.in_paragraph = False

    def _render_table(self):
        if not self.table_data:
            return
        rows = len(self.table_data)
        if rows == 0:
            return
        cols = max(len(row) for row in self.table_data)
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, row_data in enumerate(self.table_data):
            for j, cell_text in enumerate(row_data):
                if j < cols:
                    cell = table.cell(i, j)
                    cell.text = str(cell_text)
                    if i == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

    def _add_hyperlink(self, paragraph, text, url):
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = __import__("docx.oxml", fromlist=["OxmlElement"]).OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = __import__("docx.oxml", fromlist=["OxmlElement"]).OxmlElement("w:r")
        rPr = __import__("docx.oxml", fromlist=["OxmlElement"]).OxmlElement("w:rPr")
        c = __import__("docx.oxml", fromlist=["OxmlElement"]).OxmlElement("w:color")
        c.set(qn("w:val"), "0563C1")
        rPr.append(c)
        u = __import__("docx.oxml", fromlist=["OxmlElement"]).OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)


class DOCXGenerator(BaseGenerator):
    @property
    def format(self) -> str:
        return "docx"

    @property
    def mime_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    @property
    def file_extension(self) -> str:
        return ".docx"

    def generate(self, html: str, template: Optional[Dict[str, Any]] = None, metadata: Optional[Dict[str, Any]] = None) -> bytes:
        doc = Document()
        self._apply_template_styles(doc, template)
        parser = _DocxHTMLParser(doc)
        parser.feed(html)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_preview(self, html: str, template: Optional[Dict[str, Any]] = None) -> bytes:
        return self.generate(html, template)

    def _apply_template_styles(self, doc: Document, template: Optional[Dict] = None):
        if not template:
            return
        fonts = template.get("fonts", {})
        colors = template.get("colors", {})
        margins = template.get("margins", {})

        for section in doc.sections:
            if margins:
                section.top_margin = Inches(margins.get("top", 72) / 72)
                section.bottom_margin = Inches(margins.get("bottom", 72) / 72)
                section.left_margin = Inches(margins.get("left", 72) / 72)
                section.right_margin = Inches(margins.get("right", 72) / 72)


def docx_generate(html: str, template: Optional[Dict] = None, metadata: Optional[Dict] = None) -> bytes:
    gen = DOCXGenerator()
    return gen.generate(html, template, metadata)
